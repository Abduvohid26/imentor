"""Xodimlar ro'yxati (.docx) fayldan faqat o'qituvchilarni import qiladi.

Manba fayl — butun institutning xodimlar ro'yxati (bitta jadval: №, F.I.Sh
(x2), Lavozimi, Tug'ilgan yili), bo'limlar jadval ichida sarlavha-qator
sifatida keladi. Bu FAYL DARS JADVALI EMAS — faqat kim qaysi kafedrada,
qanday lavozimda ishlashini beradi.

Filtrlash:
  1. Faqat kafedra bo'limlari (ma'muriy bo'limlar — Rahbariyat, Kadrlar,
     Xisobxona va h.k. — o'tkazib yuboriladi). Diapazon: boshlanish bo'limi
     nomida "DAVOLASH ISHI" + "DEKANAT" birga uchraganda boshlanadi, "MALAKA
     OSHIRISH" bo'limida (va undan keyin) tugaydi.
  2. Har bir kafedra ichida — faqat o'qituvchi lavozimlari (Kafedra mudiri,
     Professor, Dotsent, Katta o'qituvchi, O'qituvchi, Assistent,
     Stajyor-o'qituvchi, Dekan). Kabinet mudiri, Laborant, Kotiba/Ish
     yurituvchi, Tyutor — o'qituvchi emas, o'tkazib yuboriladi.
     Lavozim yozilishi juda notekis (imloviy xato ko'p) — shuning uchun
     normallashtirilgan (apostrof/tinish belgilari olib tashlangan, kichik
     harf) kalit so'z moslashtirish ishlatiladi. Hech qaysi ro'yxatga mos
     kelmagan lavozimlar avtomatik hal qilinmaydi — "NOANIQ" ro'yxatida
     alohida ko'rsatiladi, import qilinmaydi.

Kafedra bo'limi nomi OnlineTest'dan sinxronlangan `AcademicDepartment.name`
bilan (katta-kichik harfga sezgir emas) solishtiriladi — mos kelsa
`StaffProfile.department` shu kanonik nomga o'rnatiladi, aks holda bo'lim
nomi xom holicha yoziladi va "MOS KELMAGAN KAFEDRA" ro'yxatida ko'rsatiladi.

Login: faylda telefon raqami yo'q — yangi xodim uchun vaqtinchalik
`998900000XXX` (ketma-ket) telefon va standart parol beriladi. Mavjud
xodim (F.I.Sh + kafedra orqali taxminiy moslashtirilgan) qayta yaratilmaydi,
faqat department/job_title yangilanadi.

  python manage.py import_teachers_from_docx --file "data/Mart 2026.docx"
  python manage.py import_teachers_from_docx --file "data/Mart 2026.docx" --apply
"""
from __future__ import annotations

import re

from django.contrib.auth.models import Group, User
from django.core.management.base import BaseCommand, CommandError
from django.db import transaction

from core.models import AcademicDepartment, StaffProfile

START_MARKERS = ("davolash ishi", "dekanat")
END_MARKER = "malaka oshirish"

APOSTROPHES = "'’‘‛´`"

TEACHING_KEYWORDS = [
    "assist",
    "stajyor",
    "stajor",
    "oqituvchi",
    "dotsent",
    "professor",
    "kafedra mudir",
    "dekan",
]
EXCLUDE_KEYWORDS = [
    "kabinet",
    "laborant",
    "laborator",
    "lab mud",
    "kotiba",
    "ish yurituvchi",
    "tyutor",
]

DEFAULT_PHONE_PREFIX = "998900000"
DEFAULT_PASSWORD = "TeacherTemp2026!"


def normalize(s: str) -> str:
    s = (s or "").lower()
    for ch in APOSTROPHES:
        s = s.replace(ch, "")
    s = re.sub(r"[^a-z0-9а-яё\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def classify_position(raw_position: str) -> str:
    """'teacher' | 'excluded' | 'unmatched'."""
    norm = normalize(raw_position)
    if not norm:
        return "unmatched"
    if any(k in norm for k in EXCLUDE_KEYWORDS):
        return "excluded"
    if any(k in norm for k in TEACHING_KEYWORDS):
        return "teacher"
    return "unmatched"


def split_person_name(full_name: str) -> tuple[str, str]:
    parts = [p for p in str(full_name or "").strip().split() if p]
    if not parts:
        return ("", "")
    if len(parts) == 1:
        return (parts[0], "")
    return (parts[0], " ".join(parts[1:]))


def extract_section_name(header_cell_text: str) -> str:
    """'"14 ta\\tAKUSHERLIK VA GINEKOLOGIYA"' -> 'AKUSHERLIK VA GINEKOLOGIYA'."""
    text = header_cell_text.replace("\t", " ")
    text = re.sub(r"^\s*\d+\s*ta\b", "", text, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", text).strip()


class Command(BaseCommand):
    help = "Xodimlar ro'yxati (.docx) fayldan faqat o'qituvchilarni ajratib import qiladi."

    def add_arguments(self, parser):
        parser.add_argument("--file", required=True, help="Xodimlar ro'yxati .docx fayl yo'li")
        parser.add_argument("--apply", action="store_true", help="Haqiqatan saqlash (default: dry-run)")
        parser.add_argument(
            "--default-password",
            default=DEFAULT_PASSWORD,
            help=f"Yangi xodimlar uchun vaqtinchalik parol (default: {DEFAULT_PASSWORD})",
        )

    def handle(self, *args, **opts):
        try:
            import docx
        except ImportError:
            raise CommandError("python-docx o'rnatilmagan: pip install python-docx")

        file_path = opts["file"]
        apply_changes = bool(opts["apply"])
        default_password = opts["default_password"]

        try:
            document = docx.Document(file_path)
        except Exception as exc:
            raise CommandError(f"Fayl o'qilmadi: {exc}")
        if not document.tables:
            raise CommandError("Faylda jadval topilmadi.")
        table = document.tables[0]

        # Kanonik kafedra nomlari (katta-kichik harfga sezgir emas moslashtirish uchun)
        dept_by_norm = {normalize(d.name): d.name for d in AcademicDepartment.objects.all()}

        # Mavjud xodimlar — F.I.Sh (normallashtirilgan) orqali taxminiy moslashtirish.
        # StaffProfile.owner_key oddiy CharField (User'ga FK emas) — username orqali qo'lda bog'laymiz.
        profile_owner_keys = set(StaffProfile.objects.values_list("owner_key", flat=True))
        existing_by_name: dict[str, str] = {}
        for u in User.objects.filter(username__in=profile_owner_keys):
            key = normalize(f"{u.first_name} {u.last_name}")
            if key:
                existing_by_name[key] = u.username

        # Mavjud placeholder telefonlardan keyingi bo'sh raqamni topish.
        existing_placeholders = User.objects.filter(
            username__startswith=DEFAULT_PHONE_PREFIX
        ).values_list("username", flat=True)
        next_suffix = 1
        for uname in existing_placeholders:
            suffix = uname[len(DEFAULT_PHONE_PREFIX):]
            if suffix.isdigit():
                next_suffix = max(next_suffix, int(suffix) + 1)

        in_range = False
        current_section = ""
        created = 0
        updated = 0
        excluded_count = 0
        unmatched: list[tuple[str, str, str]] = []
        unmatched_depts: set[str] = set()
        processed = 0

        with transaction.atomic():
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 4:
                    continue

                if cells[0]:
                    section = extract_section_name(cells[0])
                    norm_section = normalize(section)
                    if all(m in norm_section for m in START_MARKERS):
                        in_range = True
                    elif END_MARKER in norm_section:
                        in_range = False
                    current_section = section
                    continue

                if not in_range:
                    continue

                full_name = cells[1]
                position = cells[3]
                if not full_name or not position:
                    continue

                cls = classify_position(position)
                if cls == "excluded":
                    excluded_count += 1
                    continue
                if cls == "unmatched":
                    unmatched.append((current_section, full_name, position))
                    continue

                processed += 1
                name_key = normalize(full_name)
                dept_norm = normalize(current_section)
                dept_name = dept_by_norm.get(dept_norm)
                if not dept_name:
                    dept_name = current_section
                    unmatched_depts.add(current_section)

                existing_username = existing_by_name.get(name_key)
                if existing_username:
                    self.stdout.write(f"  [YANGILASH] {full_name} ({position}) -> {dept_name}")
                    updated += 1
                    if apply_changes:
                        profile, _ = StaffProfile.objects.get_or_create(owner_key=existing_username)
                        profile.department = dept_name
                        profile.job_title = position
                        profile.save(update_fields=["department", "job_title", "updated_at"])
                    continue

                username = f"{DEFAULT_PHONE_PREFIX}{next_suffix:03d}"
                next_suffix += 1
                first_name, last_name = split_person_name(full_name)
                self.stdout.write(f"  [YANGI] {full_name} ({position}) -> {dept_name}  login={username}")
                created += 1
                if apply_changes:
                    user = User.objects.create(username=username, first_name=first_name, last_name=last_name)
                    user.set_password(default_password)
                    user.save(update_fields=["password"])
                    group, _ = Group.objects.get_or_create(name="hodim")
                    user.groups.add(group)
                    StaffProfile.objects.create(
                        owner_key=username,
                        department=dept_name,
                        job_title=position,
                    )
                    existing_by_name[name_key] = username

            if not apply_changes:
                transaction.set_rollback(True)

        self.stdout.write("")
        self.stdout.write(
            f"NATIJA: yangi={created}  yangilandi={updated}  o'tkazib_yuborildi={excluded_count}  "
            f"noaniq={len(unmatched)}  jarayondan_o'tgan={processed}  APPLY={apply_changes}"
        )
        if unmatched_depts:
            self.stdout.write(self.style.WARNING(f"Mos kelmagan kafedra nomlari (AcademicDepartment'da topilmadi): {sorted(unmatched_depts)}"))
        if unmatched:
            self.stdout.write(self.style.WARNING("NOANIQ lavozimlar (qo'lda tekshirish kerak):"))
            for section, name, position in unmatched[:30]:
                self.stdout.write(f"    [{section}] {name} — {position!r}")
            if len(unmatched) > 30:
                self.stdout.write(f"    ... yana {len(unmatched) - 30} ta")
        if not apply_changes:
            self.stdout.write(self.style.WARNING("Bu DRY-RUN — saqlash uchun --apply qo'shing"))
